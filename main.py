import sys

from PyQt6.QtGui import QPixmap
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QPushButton, QLabel,
    QLineEdit, QComboBox, QStackedWidget, QMessageBox, QTableWidget, QTableWidgetItem, QHBoxLayout, QInputDialog,
    QFileDialog, QDialog, QAbstractScrollArea, QHeaderView
)
from PyQt6.QtCore import Qt
import sqlite3

from openpyxl import Workbook
from docx import Document
from reportlab.pdfgen import canvas

def load_stylesheet():
    with open('style.css', 'r', encoding='utf-8') as f:
        return f.read()


# База данных
class Database:
    def __init__(self, db_name="cars.db"):
        self.conn = sqlite3.connect(db_name)
        self.create_tables()

    def create_tables(self):
        with self.conn:
            self.conn.execute("""
                CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT UNIQUE,
                    password TEXT,
                    role TEXT
                )
            """)
            self.conn.execute("""
                CREATE TABLE IF NOT EXISTS cars (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    brand TEXT,
                    model TEXT,
                    year INTEGER,
                    price REAL,
                    description TEXT,
                    owner_id INTEGER,
                    FOREIGN KEY (owner_id) REFERENCES users(id)
                )
            """)
            self.conn.execute("""
                CREATE TABLE IF NOT EXISTS photos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    car_id INTEGER,
                    filepath TEXT,
                    FOREIGN KEY (car_id) REFERENCES cars(id)
                )
            """)
            self.conn.execute("""
                CREATE TABLE IF NOT EXISTS sales (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    car_id INTEGER,
                    buyer_id INTEGER,
                    FOREIGN KEY (car_id) REFERENCES cars(id),
                    FOREIGN KEY (buyer_id) REFERENCES users(id)
                )
            """)

    def register_user(self, username, password, role):
        try:
            with self.conn:
                self.conn.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)",
                                  (username, password, role))
            return True
        except sqlite3.IntegrityError:
            return False

    def authenticate_user(self, username, password):
        return self.conn.execute("SELECT id, role FROM users WHERE username = ? AND password = ?", (username, password)).fetchone()

    def add_car(self, brand, model, year, price, description, owner_id):
        with self.conn:
            cursor = self.conn.execute("INSERT INTO cars (brand, model, year, price, description, owner_id) VALUES (?, ?, ?, ?, ?, ?)", (brand, model, year, price, description, owner_id))
            return cursor.lastrowid

    def add_photo(self, car_id, filepath):
        with self.conn:
            self.conn.execute("INSERT INTO photos (car_id, filepath) VALUES (?, ?)", (car_id, filepath))

    def get_cars(self, exclude_owner=None):
        query = "SELECT id, brand, model, year, price, description, owner_id FROM cars"
        params = ()
        if exclude_owner is not None:
            query += " WHERE owner_id != ?"
            params = (exclude_owner,)
        return self.conn.execute(query, params).fetchall()

    def get_user_cars(self, owner_id):
        return self.conn.execute("SELECT id, brand, model, year, price, description FROM cars WHERE owner_id = ?",
                                 (owner_id,)).fetchall()

    def get_car_photos(self, car_id):
        return self.conn.execute("SELECT filepath FROM photos WHERE car_id = ?", (car_id,)).fetchall()

    def delete_car(self, car_id):
        with self.conn:
            self.conn.execute("DELETE FROM photos WHERE car_id = ?", (car_id,))
            self.conn.execute("DELETE FROM cars WHERE id = ?", (car_id,))

    def update_car(self, car_id, brand, model, year, price, description):
        with self.conn:
            self.conn.execute("""
                   UPDATE cars SET brand = ?, model = ?, year = ?, price = ?, description = ?
                   WHERE id = ?
               """, (brand, model, year, price, description, car_id))

    def buy_car(self, car_id, buyer_id):
        with self.conn:
            self.conn.execute("INSERT INTO sales (car_id, buyer_id) VALUES (?, ?)", (car_id, buyer_id))
            self.conn.execute("DELETE FROM cars WHERE id = ?", (car_id,))

    # def get_purchase_history(self, buyer_id):
    #     return self.conn.execute("""
    #         SELECT cars.name, cars.price FROM sales
    #         JOIN cars ON sales.car_id = cars.id
    #         WHERE sales.buyer_id = ?
    #     """, (buyer_id,)).fetchall()

# Главное окно
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Продажа автомобилей")
        self.setGeometry(300, 200, 800, 600)

        self.db = Database()
        self.stacked_widget = QStackedWidget()
        self.setCentralWidget(self.stacked_widget)

        self.login_window = LoginWindow(self.db, self)
        self.registration_window = RegistrationWindow(self.db, self)
        self.dashboard = Dashboard(self.db, self)

        self.stacked_widget.addWidget(self.login_window)
        self.stacked_widget.addWidget(self.registration_window)
        self.stacked_widget.addWidget(self.dashboard)

        self.show_login()

    def show_login(self):
        self.stacked_widget.setCurrentWidget(self.login_window)

    def show_registration(self):
        self.stacked_widget.setCurrentWidget(self.registration_window)

    def show_dashboard(self, user_id, role):
        self.dashboard.set_user(user_id, role)
        self.stacked_widget.setCurrentWidget(self.dashboard)

# Окно авторизации
class LoginWindow(QWidget):
    def __init__(self, db, main_window):
        super().__init__()
        self.db = db
        self.main_window = main_window

        layout = QVBoxLayout()



        self.username_input = QLineEdit(self)
        self.username_input.setPlaceholderText("Имя пользователя")
        self.username_input.setObjectName("Username_input")

        self.password_input = QLineEdit(self)
        self.password_input.setPlaceholderText("Пароль")
        self.password_input.setObjectName("Password_input")
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)

        self.login_button = QPushButton("Войти")
        self.login_button.setObjectName("Login_button")
        self.login_button.clicked.connect(self.login)
        self.register_button = QPushButton("Регистрация")
        self.register_button.setObjectName("register_button")
        self.register_button.clicked.connect(self.main_window.show_registration)

        layout.addWidget(QLabel("Авторизация"))
        layout.addWidget(self.username_input)
        layout.addWidget(self.password_input)
        layout.addWidget(self.login_button)
        layout.addWidget(self.register_button)

        self.setLayout(layout)
        self.setStyleSheet("""
                    QLineEdit#Username_input, QLineEdit#Password_input {
                        border: 2px solid blue;
                        border-radius: 5px;
                        padding: 5px;
                    }
                    QPushButton#register_button, QPushButton#Login_button{
                        background-color: blue;
                        color: white;
                        border: none;
                        border-radius: 5px;
                        padding: 10px;
                    }
                    QPushButton#Login_button:hover {
                        background-color: blue;
                    }
                """)

    def login(self):
        username = self.username_input.text()
        password = self.password_input.text()
        if not username or not password:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, заполните все поля")
            return
        user = self.db.authenticate_user(username, password)

        if user:
            user_id, role = user
            self.main_window.show_dashboard(user_id, role)
        else:
            QMessageBox.warning(self, "Ошибка", "Неверное имя пользователя или пароль")


# Окно регистрации
class RegistrationWindow(QWidget):
    def __init__(self, db, main_window):
        super().__init__()
        self.db = db
        self.main_window = main_window


        layout = QVBoxLayout()

        self.username_input = QLineEdit(self)
        self.username_input.setPlaceholderText("Имя пользователя")
        self.username_input.setObjectName("usernameInput")

        self.password_input = QLineEdit(self)
        self.password_input.setPlaceholderText("Пароль")
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.password_input.setObjectName("passwordInput")

        self.role_input = QComboBox(self)
        self.role_input.addItems(["Покупатель", "Продавец"])
        self.role_input.setObjectName("roleInput")

        self.register_button = QPushButton("Зарегистрироваться")
        self.register_button.setObjectName("registerButton")
        self.register_button.clicked.connect(self.register)

        layout.addWidget(QLabel("Регистрация"))
        layout.addWidget(self.username_input)
        layout.addWidget(self.password_input)
        layout.addWidget(self.role_input)
        layout.addWidget(self.register_button)

        self.setLayout(layout)
        self.setStyleSheet("""
            QLineEdit#usernameInput, QLineEdit#passwordInput {
                border: 2px solid blue;
                border-radius: 5px;
                padding: 5px;
            }
            QPushButton#registerButton {
                background-color: blue;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 10px;
            }
            QPushButton#registerButton:hover {
                background-color: blue;
            }
        """)

    def register(self):
        username = self.username_input.text()
        password = self.password_input.text()
        role = self.role_input.currentText()
        if not username or not password:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, заполните все поля")
            return

        if self.db.register_user(username, password, role):
            QMessageBox.information(self, "Успех", "Регистрация прошла успешно")
            self.main_window.show_login()
        else:
            QMessageBox.warning(self, "Ошибка", "Имя пользователя уже занято")

# Панель управления
class Dashboard(QWidget):
    def __init__(self, db, main_window):
        super().__init__()
        self.db = db
        self.main_window = main_window
        self.user_id = None
        self.role = None
        self.theme = "light"

        self.layout = QVBoxLayout()
        self.setStyleSheet(load_stylesheet())  # Применяем стили

        self.user_label = QLabel("Добро пожаловать!")
        self.logout_button = QPushButton("Выйти")
        self.logout_button.clicked.connect(self.logout)
        self.theme_button = QPushButton("Сменить тему")
        self.theme_button.clicked.connect(self.toggle_theme)

        self.layout.addWidget(self.user_label)
        self.layout.addWidget(self.theme_button)
        self.layout.addWidget(self.logout_button)
        self.setLayout(self.layout)

    def set_user(self, user_id, role):
        self.user_id = user_id
        self.role = role
        self.user_label.setText(f"Добро пожаловать! Роль: {role}")
        self.update_dashboard()

    def update_dashboard(self):
        for i in reversed(range(self.layout.count())):
            widget = self.layout.itemAt(i).widget()
            if widget:
                widget.deleteLater()

        self.layout.addWidget(self.user_label)
        self.layout.addWidget(self.theme_button)
        self.layout.addWidget(self.logout_button)

        if self.role == "Продавец":
            self.add_seller_dashboard()
        elif self.role == "Покупатель":
            self.add_buyer_dashboard()

    def add_seller_dashboard(self):
        add_car_button = QPushButton("Добавить автомобиль")
        add_car_button.clicked.connect(self.add_car)
        self.layout.addWidget(add_car_button)

        self.my_cars_table = QTableWidget()
        self.update_my_cars_table()
        self.layout.addWidget(self.my_cars_table)

    def add_buyer_dashboard(self):
        # Поле поиска
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Введите марку или модель для поиска...")
        self.search_input.textChanged.connect(self.update_available_cars_table)
        self.layout.addWidget(self.search_input)

        # Таблица доступных автомобилей
        self.cars_table = QTableWidget()
        self.update_available_cars_table()
        self.layout.addWidget(self.cars_table)

        # Кнопка истории покупок
        view_history_button = QPushButton("История покупок")
        view_history_button.clicked.connect(self.view_purchase_history)
        self.layout.addWidget(view_history_button)

    def update_my_cars_table(self):
        cars = self.db.get_user_cars(self.user_id)
        self.my_cars_table.setRowCount(len(cars))
        self.my_cars_table.setColumnCount(8)
        header = self.my_cars_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(6, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(7, QHeaderView.ResizeMode.ResizeToContents)

        self.my_cars_table.setHorizontalHeaderLabels(
            ["Марка", "Модель", "Год", "Цена", "Описание", "Фотографии", "Добавление", "Удаление"]
        )

        for row, (car_id, brand, model, year, price, description) in enumerate(cars):
            self.my_cars_table.setItem(row, 0, QTableWidgetItem(brand))
            self.my_cars_table.setItem(row, 1, QTableWidgetItem(model))
            self.my_cars_table.setItem(row, 2, QTableWidgetItem(str(year)))
            self.my_cars_table.setItem(row, 3, QTableWidgetItem(str(price)))
            self.my_cars_table.setItem(row, 4, QTableWidgetItem(description))

            # Фотографии: добавляем кнопку для просмотра
            photos_button = QPushButton("Просмотреть")
            photos_button.clicked.connect(lambda checked, car_id=car_id: self.view_photos(car_id))

            # Действия: кнопки редактирования и удаления

            edit_button = QPushButton("Изменить")
            edit_button.setObjectName("editButton")
            edit_button.clicked.connect(lambda checked, car_id=car_id: self.edit_car(car_id))


            delete_button = QPushButton("Удалить")
            delete_button.setObjectName("deleteButton")
            delete_button.clicked.connect(lambda checked, car_id=car_id: self.delete_car(car_id))

            self.my_cars_table.setCellWidget(row, 5, photos_button)  # Фото
            self.my_cars_table.setCellWidget(row, 6, edit_button)  # Действия
            self.my_cars_table.setCellWidget(row, 7, delete_button)

        self.my_cars_table.resizeColumnsToContents()

    def view_photos(self, car_id):
        photos = self.db.get_car_photos(car_id)
        if not photos:
            QMessageBox.information(self, "Нет фотографий", "У этого автомобиля нет фотографий.")
            return

        current_index = 0

        def show_photo(index):
            pixmap = QPixmap(photos[index][0])
            photo_label.setPixmap(pixmap.scaled(400, 400, Qt.AspectRatioMode.KeepAspectRatio))

        photo_dialog = QDialog(self)
        photo_layout = QVBoxLayout()
        photo_label = QLabel()
        photo_layout.addWidget(photo_label)

        prev_button = QPushButton("Назад")
        next_button = QPushButton("Вперед")

        def prev_photo():
            nonlocal current_index
            current_index = (current_index - 1) % len(photos)
            show_photo(current_index)

        def next_photo():
            nonlocal current_index
            current_index = (current_index + 1) % len(photos)
            show_photo(current_index)

        prev_button.clicked.connect(prev_photo)
        next_button.clicked.connect(next_photo)

        button_layout = QHBoxLayout()
        button_layout.addWidget(prev_button)
        button_layout.addWidget(next_button)

        photo_layout.addLayout(button_layout)
        photo_dialog.setLayout(photo_layout)
        photo_dialog.setWindowTitle("Фотографии автомобиля")
        show_photo(current_index)
        photo_dialog.exec()

    def update_available_cars_table(self):
        # Получение автомобилей из базы
        cars = self.db.get_cars(exclude_owner=self.user_id)

        # Фильтрация по поисковому запросу
        search_text = self.search_input.text().lower() if hasattr(self, 'search_input') else ""
        if search_text:
            cars = [car for car in cars if search_text in car[1].lower() or search_text in car[2].lower()]

        self.cars_table.setRowCount(len(cars))
        self.cars_table.setColumnCount(4)
        self.cars_table.setHorizontalHeaderLabels(["Марка", "Модель", "Цена", "Купить"])

        for row, (car_id, brand, model, year, price, description, owner_id) in enumerate(cars):
            self.cars_table.setItem(row, 0, QTableWidgetItem(brand))
            self.cars_table.setItem(row, 1, QTableWidgetItem(model))
            self.cars_table.setItem(row, 2, QTableWidgetItem(str(price)))

            buy_button = QPushButton("Купить")
            buy_button.clicked.connect(lambda checked, car_id=car_id: self.buy_car(car_id))
            self.cars_table.setCellWidget(row, 3, buy_button)

    def add_car(self):
        brand, ok = QInputDialog.getText(self, "Добавить автомобиль", "Введите марку автомобиля:")
        if ok and brand:
            model, ok = QInputDialog.getText(self, "Добавить автомобиль", "Введите модель автомобиля:")
            if ok and model:
                year, ok = QInputDialog.getInt(self, "Добавить автомобиль", "Введите год выпуска:")
                if ok:
                    price, ok = QInputDialog.getDouble(self, "Добавить автомобиль", "Введите цену:")
                    if ok:
                        description, ok = QInputDialog.getText(self, "Добавить автомобиль", "Введите описание:")
                        if ok:
                            car_id = self.db.add_car(brand, model, year, price, description, self.user_id)
                            self.add_car_photos(car_id)
                            self.update_my_cars_table()

    def add_car_photos(self, car_id):
        file_dialog = QFileDialog(self)
        file_dialog.setFileMode(QFileDialog.FileMode.ExistingFiles)
        file_dialog.setNameFilter("Изображения (*.png *.jpg *.jpeg)")
        if file_dialog.exec():
            file_paths = file_dialog.selectedFiles()
            for file_path in file_paths:
                self.db.add_photo(car_id, file_path)

    def edit_car(self, car_id):
        brand, ok = QInputDialog.getText(self, "Изменить автомобиль", "Введите новую марку:")
        if ok and brand:
            model, ok = QInputDialog.getText(self, "Изменить автомобиль", "Введите новую модель:")
            if ok and model:
                year, ok = QInputDialog.getInt(self, "Изменить автомобиль", "Введите новый год выпуска:")
                if ok:
                    price, ok = QInputDialog.getDouble(self, "Изменить автомобиль", "Введите новую цену:")
                    if ok:
                        description, ok = QInputDialog.getText(self, "Изменить автомобиль", "Введите новое описание:")
                        if ok:
                            self.db.update_car(car_id, brand, model, year, price, description)
                            self.update_my_cars_table()

    def export_to_excel(self, cars):
        wb = Workbook()
        ws = wb.active
        ws.append(["Марка", "Модель", "Год", "Цена", "Описание"])
        for car in cars:
            ws.append(car[1:6])  # Игнорируем ID и владелец
        filename, _ = QFileDialog.getSaveFileName(None, "Сохранить как", "", "Excel Files (*.xlsx)")
        if filename:
            wb.save(filename)
            QMessageBox.information(None, "Успех", f"Данные экспортированы в {filename}.")

    def export_to_word(self, cars):
        doc = Document()
        doc.add_heading("Список автомобилей", level=1)
        for car in cars:
            doc.add_paragraph(f"Марка: {car[1]}, Модель: {car[2]}, Год: {car[3]}, Цена: {car[4]}, Описание: {car[5]}")
        filename, _ = QFileDialog.getSaveFileName(None, "Сохранить как", "", "Word Files (*.docx)")
        if filename:
            doc.save(filename)
            QMessageBox.information(None, "Успех", f"Данные экспортированы в {filename}.")

    # def export_to_pdf(self, cars):
    #     filename, _ = QFileDialog.getSaveFileName(None, "Сохранить как", "", "PDF Files (*.pdf)")
    #     if filename:
    #         c = canvas.Canvas(filename)
    #         c.setFont("Times-Roman", 12)
    #         y = 800
    #         for car in cars:
    #             c.drawString(50, y,
    #                          f"Марка: {car[1]}, Модель: {car[2]}, Год: {car[3]}, Цена: {car[4]}, Описание: {car[5]}"
    #                          )
    #             y -= 20
    #             if y < 50:
    #                 c.showPage()
    #                 y = 800

    #         c.save()
    #         QMessageBox.information(None, "Успех", f"Данные экспортированы в {filename}.")

    def add_seller_dashboard(self):
        export_layout = QHBoxLayout()
        export_excel_button = QPushButton("Экспорт в Excel")
        export_excel_button.clicked.connect(lambda: self.export_to_excel(self.db.get_user_cars(self.user_id)))

        export_word_button = QPushButton("Экспорт в Word")
        export_word_button.clicked.connect(lambda: self.export_to_word(self.db.get_user_cars(self.user_id)))

        # export_pdf_button = QPushButton("Экспорт в PDF")
        # export_pdf_button.clicked.connect(lambda: self.export_to_pdf(self.db.get_user_cars(self.user_id)))

        export_layout.addWidget(export_excel_button)
        export_layout.addWidget(export_word_button)
        # export_layout.addWidget(export_pdf_button)
        self.layout.addLayout(export_layout)

        add_car_button = QPushButton("Добавить автомобиль")
        add_car_button.clicked.connect(self.add_car)
        self.layout.addWidget(add_car_button)

        self.my_cars_table = QTableWidget()
        self.update_my_cars_table()
        self.layout.addWidget(self.my_cars_table)


    def delete_car(self, car_id):
        self.db.delete_car(car_id)
        self.update_my_cars_table()

    def buy_car(self, car_id):
        self.db.buy_car(car_id, self.user_id)
        self.update_available_cars_table()

    def view_purchase_history(self):
         history = self.db.conn.execute("""
             SELECT cars.brand, cars.model, cars.year, sales.car_id
             FROM sales
            LEFT JOIN cars ON sales.car_id = cars.id
             WHERE sales.buyer_id = ?
         """, (self.user_id,)).fetchall()

         if not history:
             QMessageBox.information(self, "История покупок", "Вы еще ничего не купили.")
             return

         history_text = "\n".join([
             f"ID: {car_id}, {brand} {model} ({year})"
             for brand, model, year, car_id in history
         ])
         QMessageBox.information(self, "История покупок", history_text)

    def toggle_theme(self):
        if self.theme == "light":
            QApplication.instance().setStyleSheet("QWidget { background-color: #121212; color: #ffffff; }")
            self.theme = "dark"
        else:
            QApplication.instance().setStyleSheet("")
            self.theme = "light"

    def logout(self):
        self.main_window.show_login()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    main_window = MainWindow()
    main_window.show()
    sys.exit(app.exec())
